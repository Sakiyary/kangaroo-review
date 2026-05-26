#!/usr/bin/env python3
import argparse
import hashlib
import json
import re
import subprocess
from pathlib import Path

from docx import Document
from PIL import Image
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]


def slugify(path: Path) -> str:
    raw = path.as_posix()
    digest = hashlib.sha1(raw.encode("utf-8")).hexdigest()[:8]
    stem = re.sub(r"[^A-Za-z0-9]+", "-", path.stem).strip("-").lower()
    return f"{stem or 'source'}-{digest}"


def read_markdown(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def read_rtf(path: Path) -> str:
    result = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        check=True,
        capture_output=True,
        text=True,
    )
    return result.stdout


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        if rows:
            parts.append("\n".join(rows))
    return "\n\n".join(parts)


def read_pdf(path: Path) -> tuple[list[dict], int, int]:
    reader = PdfReader(str(path))
    pages = []
    chars = 0
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        chars += len(text)
        pages.append({"page": index, "chars": len(text), "text": text})
    return pages, len(reader.pages), chars


def image_meta(path: Path) -> dict:
    with Image.open(path) as image:
        return {"width": image.width, "height": image.height, "mode": image.mode}


def write_text(out_dir: Path, source: Path, text: str) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{slugify(source)}.md"
    rel = source.relative_to(ROOT)
    out_path.write_text(f"# {rel.as_posix()}\n\n{text}\n", encoding="utf-8")
    return out_path


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--roots", nargs="+", default=["slides", "raw"])
    parser.add_argument("--out", default="data")
    args = parser.parse_args()

    out_root = ROOT / args.out
    catalog_dir = out_root / "catalog"
    catalog_dir.mkdir(parents=True, exist_ok=True)

    source_records = []
    for root_name in args.roots:
        root = ROOT / root_name
        if not root.exists():
            continue
        for path in sorted(p for p in root.rglob("*") if p.is_file() and p.name != ".DS_Store"):
            rel = path.relative_to(ROOT)
            suffix = path.suffix.lower()
            record = {
                "path": rel.as_posix(),
                "bytes": path.stat().st_size,
                "kind": suffix.lstrip(".") or "unknown",
                "extraction": "not_attempted",
                "text_chars": 0,
            }
            try:
                if suffix in {".md", ".txt"}:
                    text = read_markdown(path)
                    target = write_text(out_root / "extracted" / root_name, path, text)
                    record.update(
                        extraction="text",
                        text_chars=len(text),
                        extracted_path=target.relative_to(ROOT).as_posix(),
                    )
                elif suffix == ".rtf":
                    text = read_rtf(path)
                    target = write_text(out_root / "extracted" / root_name, path, text)
                    record.update(
                        extraction="textutil",
                        text_chars=len(text),
                        extracted_path=target.relative_to(ROOT).as_posix(),
                    )
                elif suffix == ".docx":
                    text = read_docx(path)
                    target = write_text(out_root / "extracted" / root_name, path, text)
                    record.update(
                        extraction="python-docx",
                        text_chars=len(text),
                        extracted_path=target.relative_to(ROOT).as_posix(),
                    )
                elif suffix == ".pdf":
                    pages, page_count, chars = read_pdf(path)
                    target_dir = out_root / "extracted" / root_name
                    target_dir.mkdir(parents=True, exist_ok=True)
                    target = target_dir / f"{slugify(path)}.json"
                    target.write_text(
                        json.dumps(
                            {
                                "source": rel.as_posix(),
                                "page_count": page_count,
                                "text_chars": chars,
                                "pages": pages,
                            },
                            ensure_ascii=False,
                            indent=2,
                        ),
                        encoding="utf-8",
                    )
                    record.update(
                        extraction="pypdf",
                        page_count=page_count,
                        text_chars=chars,
                        extracted_path=target.relative_to(ROOT).as_posix(),
                        needs_ocr=chars < max(200, page_count * 80),
                    )
                elif suffix in {".png", ".jpg", ".jpeg", ".webp"}:
                    record.update(extraction="image-meta", image=image_meta(path), needs_ocr=True)
                else:
                    record.update(extraction="unsupported")
            except Exception as exc:
                record.update(extraction="failed", error=str(exc))
            source_records.append(record)

    catalog_path = catalog_dir / "sources.json"
    catalog_path.write_text(json.dumps(source_records, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {catalog_path.relative_to(ROOT)} with {len(source_records)} records")


if __name__ == "__main__":
    main()
